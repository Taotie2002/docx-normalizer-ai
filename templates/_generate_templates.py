#!/usr/bin/env python3
"""Generate DFGP template .docx files for notification, report, letter, academic."""

from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# GB/T 9704-2012 constants
MARGINS = dict(top=Cm(3.7), bottom=Cm(3.5), left=Cm(2.8), right=Cm(2.6))
PAGE_SIZE = dict(width=Cm(21), height=Cm(29.7))
LINE_SPACING = Pt(28)  # 固定值28pt
FONT_BODY = "仿宋_GB2312"
FONT_BODY_CS = "仿宋_GB2312"
FONT_TITLE = "方正小标宋简体"
FONT_H1 = "黑体"
FONT_H2 = "楷体_GB2312"
FONT_H2_CS = "楷体_GB2312"
FONT_SALUTE = "仿宋_GB2312"
FONT_SIGN = "仿宋_GB2312"
SZ_BODY = 16  # 16pt = 32 half-points
SZ_TITLE = 22
SZ_H1 = 16
SZ_H2 = 16
SZ_PAGE_NUM = 14

def set_margins(section):
    section.page_width = PAGE_SIZE["width"]
    section.page_height = PAGE_SIZE["height"]
    section.top_margin = MARGINS["top"]
    section.bottom_margin = MARGINS["bottom"]
    section.left_margin = MARGINS["left"]
    section.right_margin = MARGINS["right"]

def apply_run_font(run, font_name, font_cs, size_pt):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = False

def set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=None, line_spacing_rule=None, space_before=None, space_after=None):
    pf = para.paragraph_format
    pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if line_spacing_rule is not None:
        pf.line_spacing_rule = line_spacing_rule
        pf.line_spacing = LINE_SPACING
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after

def add_para(doc, text, font_name, font_cs, size_pt, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line_indent=None, bold=False, space_before=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    apply_run_font(run, font_name, font_cs, size_pt)
    run.font.bold = bold
    set_para_format(p, alignment=alignment,
                    first_line_indent=first_line_indent,
                    line_spacing_rule=WD_LINE_SPACING.MULTIPLE if first_line_indent else None,
                    space_before=space_before, space_after=space_after)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

def configure_page_for_gbt(doc):
    """Set document-level page settings."""
    # python-docx section approach
    section = doc.sections[0]
    set_margins(section)
    # Set titlePg (different first page) via XML
    sectPr = section._sectPr
    titlePg = OxmlElement("w:titlePg")
    titlePg.set(qn("w:val"), "1")
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is not None:
        pgMar.set(qn("w:titlePg"), "1")
    else:
        sectPr.insert(0, titlePg)
    # Set pgSz
    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.insert(0, pgSz)
    pgSz.set(qn("w:w"), str(int(PAGE_SIZE["width"].twips)))
    pgSz.set(qn("w:h"), str(int(PAGE_SIZE["height"].twips)))
    # Mirror margins
    pgMar2 = sectPr.find(qn("w:pgMar"))
    if pgMar2 is not None:
        pgMar2.set(qn("w:mirrorMargins"), "1")

# ─────────────────────────────────────────────────────────────────────────────
# 1. NOTIFICATION (通知)
# ─────────────────────────────────────────────────────────────────────────────
def create_notification():
    doc = Document()
    configure_page_for_gbt(doc)

    # 主标题
    add_para(doc, "关于开展2026年度信息安全检查的通知",
             FONT_TITLE, FONT_TITLE, SZ_TITLE,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=Pt(0), space_after=Pt(0))

    # 发文编号行（可选）
    add_para(doc, "××〔2026〕12号",
             FONT_BODY, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=Pt(0), space_after=Pt(6))

    # 称谓行
    add_para(doc, "各市（州）人民政府，省直各部门：",
             FONT_SALUTE, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=Pt(6), space_after=Pt(6))

    # 正文
    body_paras = [
        "为深入贯彻落实《网络安全法》《数据安全法》《个人信息保护法》以及国家和省关于信息安全工作的决策部署，进一步加强全省信息安全管理工作，提升网络安全防护能力，经研究，决定在全省范围内开展2026年度信息安全检查工作。现将有关事项通知如下：",

        "一、检查范围",
        "（一）检查对象：全省各级党政机关、事业单位、国有企业以及关键信息基础设施运营单位。",
        "（二）检查内容：网络安全管理制度建设情况、网络安全技术防护措施、应急响应机制、数据安全保护、个人信息保护、网络安全教育培训等方面。",

        "二、检查方式",
        "（一）自查阶段（2026年5月1日至5月31日）。各部门、各单位按照检查内容要求，认真开展自查，形成自查报告。",
        "（二）抽查阶段（2026年6月1日至6月30日）。省信息安全工作领导小组办公室将组织专家对部分单位进行实地抽查。",

        "三、工作要求",
        "（一）加强组织领导。各单位要高度重视，明确责任，确保检查工作落实到位。",
        "（二）认真自查整改。对自查发现的问题，要制定整改措施，明确整改时限。",
        "（三）按时报送材料。请于2026年5月31日前将自查报告（含电子版）报送至省信息安全工作领导小组办公室。",

        "四、联系人及方式",
        "联系人：省信息安全工作领导小组办公室  李明（联系电话：010-12345678）",
        "邮　箱：security@example.gov.cn",
    ]

    for i, text in enumerate(body_paras):
        if text.startswith("一、") or text.startswith("二、") or text.startswith("三、") or text.startswith("四、"):
            # 一级标题
            add_para(doc, text, FONT_H1, FONT_H1, SZ_H1,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Pt(0),
                     space_before=Pt(12), space_after=Pt(6))
        elif text.startswith("（一）") or text.startswith("（二）") or text.startswith("（三）"):
            # 二级标题
            add_para(doc, text, FONT_H2, FONT_H2_CS, SZ_H2,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Pt(0),
                     space_before=Pt(0), space_after=Pt(0))
        else:
            # 正文
            add_para(doc, text, FONT_BODY, FONT_BODY_CS, SZ_BODY,
                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=Pt(32),
                     space_before=Pt(0), space_after=Pt(0))

    # 结语
    add_para(doc, "附件：2026年度信息安全检查自查报告模板",
             FONT_BODY, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line_indent=Pt(32),
             space_before=Pt(12), space_after=Pt(0))

    # 落款
    add_para(doc, "省信息安全工作领导小组办公室",
             FONT_SIGN, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=Pt(24), space_after=Pt(0))
    add_para(doc, "2026年4月10日",
             FONT_SIGN, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=Pt(0), space_after=Pt(0))

    return doc

# ─────────────────────────────────────────────────────────────────────────────
# 2. REPORT (报告)
# ─────────────────────────────────────────────────────────────────────────────
def create_report():
    doc = Document()
    configure_page_for_gbt(doc)

    # 主标题
    add_para(doc, "关于2025年度全省政务公开工作情况的报告",
             FONT_TITLE, FONT_TITLE, SZ_TITLE,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=Pt(0), space_after=Pt(0))

    # 收文机关
    add_para(doc, "省政府：",
             FONT_SALUTE, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=Pt(6), space_after=Pt(6))

    # 正文
    body_paras = [
        "根据《国务院办公厅关于全面推进政务公开工作的意见》（国办发〔2026〕8号）要求，现将2025年度全省政务公开工作情况报告如下：",

        "一、基本情况",
        "2025年，全省各级行政机关认真贯彻落实党中央、国务院关于政务公开工作的决策部署，坚持以人民为中心的发展思想，扎实推进决策、执行、管理、服务、结果公开，政务公开工作取得新的成效。全年累计主动公开政府信息125.6万余条，受理政府信息公开申请2.3万件，按时办结率达98.7%。",

        "二、主要做法",
        "（一）加强制度建设，完善政务公开工作机制。出台了《全省政务公开工作考核办法》，将政务公开纳入政府绩效考核体系。",
        "（二）深化重点领域信息公开。围绕'六稳''六保'、疫情防控、乡村振兴等重点工作，主动公开各类政策文件和工作进展信息3.2万余条。",
        "（三）优化公开平台建设。完成省级政府门户网站改版升级，开设政策解读专栏，累计发布政策解读材料860余篇。",
        "（四）强化监督考核。建立健全政务公开工作年度报告制度，定期开展专项检查，对发现的问题及时督促整改。",

        "三、存在的问题",
        "（一）部分基层单位对政务公开工作重视不够，工作推进不平衡。",
        "（二）政策解读质量有待提高，解读形式较为单一。",
        "（三）政务公开专业化队伍建设仍需加强。",

        "四、下一步工作打算",
        "（一）进一步完善政务公开制度体系，修订《全省政府信息公开规定》。",
        "（二）持续深化重点领域信息公开，增强公开的针对性和时效性。",
        "（三）加强政策解读和回应关切，丰富解读形式，提升解读质量。",
        "（四）加大培训力度，提高政务公开工作人员业务能力。",

        "妥否，请批示。",
    ]

    for text in body_paras:
        if text.startswith("一、") or text.startswith("二、") or text.startswith("三、") or text.startswith("四、"):
            add_para(doc, text, FONT_H1, FONT_H1, SZ_H1,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Pt(0),
                     space_before=Pt(12), space_after=Pt(6))
        elif text.startswith("（一）") or text.startswith("（二）") or text.startswith("（三）") or text.startswith("（四）"):
            add_para(doc, text, FONT_H2, FONT_H2_CS, SZ_H2,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Pt(0),
                     space_before=Pt(0), space_after=Pt(0))
        elif "妥否" in text:
            add_para(doc, text, FONT_BODY, FONT_BODY_CS, SZ_BODY,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     first_line_indent=Pt(0),
                     space_before=Pt(12), space_after=Pt(0))
        else:
            add_para(doc, text, FONT_BODY, FONT_BODY_CS, SZ_BODY,
                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=Pt(32),
                     space_before=Pt(0), space_after=Pt(0))

    # 落款
    add_para(doc, "省政府办公厅",
             FONT_SIGN, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=Pt(24), space_after=Pt(0))
    add_para(doc, "2026年3月15日",
             FONT_SIGN, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=Pt(0), space_after=Pt(0))

    return doc

# ─────────────────────────────────────────────────────────────────────────────
# 3. LETTER (函)
# ─────────────────────────────────────────────────────────────────────────────
def create_letter():
    doc = Document()
    configure_page_for_gbt(doc)

    # 标题行（函的标题通常不加红头，直接写标题）
    add_para(doc, "关于商请协助开展联合调研的函",
             FONT_TITLE, FONT_TITLE, SZ_TITLE,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=Pt(0), space_after=Pt(0))

    # 发文字号
    add_para(doc, "××〔2026〕28号",
             FONT_BODY, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=Pt(0), space_after=Pt(6))

    # 收文机关
    add_para(doc, "省发展改革委：",
             FONT_SALUTE, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=Pt(6), space_after=Pt(6))

    # 正文
    body_paras = [
        "为深入贯彻落实党的二十届三中全会精神，加快推进全省新能源产业高质量发展，根据省政府工作部署，我委拟联合贵委开展全省新能源产业发展情况专项调研。现就有关事宜函商如下：",

        "一、调研目的",
        "全面了解全省新能源产业发展现状、存在的主要问题和困难，研究提出加快新能源产业发展的对策建议，为省委、省政府科学决策提供参考依据。",

        "二、调研内容",
        "（一）全省新能源产业发展基本情况，包括光伏、风电、生物质能等各领域装机容量和发电量；",
        "（二）新能源重点项目建设进展及存在的主要问题；",
        "（三）新能源产业技术创新和产业链发展情况；",
        "（四）新能源消纳和电网接入情况；",
        "（五）支持新能源产业发展的政策措施及落实情况。",

        "三、调研时间和方式",
        "拟于2026年5月中旬至6月上旬赴有关市（州）开展实地调研，调研方式包括座谈交流、实地考察、问卷调查等。具体调研行程另行通知。",

        "四、相关事项",
        "（一）请贵委明确1名分管领导牵头负责此项工作，并确定1名联络员；",
        "（二）请贵委协助提供新能源产业相关政策文件和统计数据；",
        "（三）调研期间有关接待工作由我委负责。",

        "专此函商，请予支持为盼。",
    ]

    for text in body_paras:
        if text.startswith("一、") or text.startswith("二、") or text.startswith("三、") or text.startswith("四、"):
            add_para(doc, text, FONT_H1, FONT_H1, SZ_H1,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Pt(0),
                     space_before=Pt(12), space_after=Pt(6))
        elif text.startswith("（一）") or text.startswith("（二）") or text.startswith("（三）") or text.startswith("（四）"):
            add_para(doc, text, FONT_H2, FONT_H2_CS, SZ_H2,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Pt(0),
                     space_before=Pt(0), space_after=Pt(0))
        elif text == "专此函商，请予支持为盼。":
            add_para(doc, text, FONT_BODY, FONT_BODY_CS, SZ_BODY,
                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=Pt(32),
                     space_before=Pt(12), space_after=Pt(0))
        else:
            add_para(doc, text, FONT_BODY, FONT_BODY_CS, SZ_BODY,
                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_line_indent=Pt(32),
                     space_before=Pt(0), space_after=Pt(0))

    # 落款
    add_para(doc, "省工业和信息化厅",
             FONT_SIGN, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=Pt(24), space_after=Pt(0))
    add_para(doc, "2026年4月8日",
             FONT_SIGN, FONT_BODY_CS, SZ_BODY,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=Pt(0), space_after=Pt(0))

    return doc

# ─────────────────────────────────────────────────────────────────────────────
# 4. ACADEMIC (学术论文)
# ─────────────────────────────────────────────────────────────────────────────
def create_academic():
    doc = Document()
    # Academic papers usually use different margins
    section = doc.sections[0]
    section.page_width = PAGE_SIZE["width"]
    section.page_height = PAGE_SIZE["height"]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    FONT_SIMHEI = "黑体"
    FONT_KAITI = "楷体_GB2312"
    FONT_SONGTI = "宋体"
    FONT_LH = "Times New Roman"

    # Title
    p = doc.add_paragraph()
    run = p.add_run("基于深度学习的政务文档智能分类方法研究")
    run.font.name = FONT_SIMHEI
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SIMHEI)
    run.font.size = Pt(18)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = Pt(24)

    # Author
    p = doc.add_paragraph()
    run = p.add_run("张伟1,2　李明1,3　王芳1")
    run.font.name = FONT_LH
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    # Affiliation
    p = doc.add_paragraph()
    run = p.add_run("1. 清华大学 计算机科学与技术系，北京 100084")
    run.font.name = FONT_KAITI
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KAITI)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("2. 中国科学院 计算技术研究所，北京 100190")
    run.font.name = FONT_KAITI
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KAITI)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("3. 华为技术有限公司，深圳 518129")
    run.font.name = FONT_KAITI
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KAITI)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    # Divider
    p = doc.add_paragraph()
    run = p.add_run("─" * 40)
    run.font.name = FONT_LH
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)

    # Abstract heading
    p = doc.add_paragraph()
    run = p.add_run("摘　要")
    run.font.name = FONT_SIMHEI
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SIMHEI)
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    # Abstract body
    p = doc.add_paragraph()
    run = p.add_run("针对政务文档自动分类准确率低、领域适应性差的问题，提出一种基于多任务学习与注意力机制的政务文档智能分类方法。该方法以BERT预训练模型为基础框架，融合文档结构特征与语义特征，构建多标签分类模型。在真实政务数据集上的实验结果表明，本文方法在_macro-F1_和_micro-F1_指标上分别达到94.3%和96.1%，较现有方法有显著提升。")
    run.font.name = FONT_KAITI
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KAITI)
    run.font.size = Pt(11)
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(6)

    # Keywords
    p = doc.add_paragraph()
    run = p.add_run("关键词：")
    run.font.name = FONT_SIMHEI
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SIMHEI)
    run.font.size = Pt(11)
    run.font.bold = True
    run2 = p.add_run("政务文档分类；深度学习；BERT；注意力机制；多任务学习")
    run2.font.name = FONT_KAITI
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KAITI)
    run2.font.size = Pt(11)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(12)

    # Section headings style helper
    def section_heading(doc, text, font_name, font_cs, size, bold=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cs)
        run.font.size = Pt(size)
        run.font.bold = bold
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def body_para(doc, text, font_name, font_cs, size):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cs)
        run.font.size = Pt(size)
        p.paragraph_format.first_line_indent = Pt(22)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    # 1 引言
    section_heading(doc, "1　引言", FONT_SIMHEI, FONT_SIMHEI, 12)
    body_para(doc, "政务文档是政府机关履行职能、处理事务的重要载体，涵盖通知、报告、函件、请示、决定等多种文体。随着数字政府建设的深入推进，海量政务文档的自动化管理与智能检索需求日益迫切。传统的基于关键词匹配和规则模板的文档分类方法，难以适应政务文档语言规范、领域专业性强、结构多样等特点，分类准确率较低。", FONT_KAITI, FONT_KAITI, 11)
    body_para(doc, "近年来，深度学习技术在自然语言处理领域取得突破性进展，BERT等预训练语言模型在多项任务上刷新了性能记录。如何将预训练模型有效迁移至政务文档分类任务，成为研究热点。", FONT_KAITI, FONT_KAITI, 11)

    # 2 相关工作
    section_heading(doc, "2　相关工作", FONT_SIMHEI, FONT_SIMHEI, 12)
    body_para(doc, "文档分类方法经历了从传统机器学习到深度学习的发展历程。早期方法主要依赖TF-IDF等特征表示，结合SVM、朴素贝叶斯等分类器。深度学习方法则通过卷积神经网络（CNN）和循环神经网络（RNN）自动学习文本表示，显著提升了分类性能。", FONT_KAITI, FONT_KAITI, 11)

    # 3 方法
    section_heading(doc, "3　本文方法", FONT_SIMHEI, FONT_SIMHEI, 12)
    section_heading(doc, "3.1　模型架构", FONT_SIMHEI, FONT_SIMHEI, 12)
    body_para(doc, "本文方法以BERT为基础编码器，融合政务文档的标题、正文、附件说明等结构化信息，采用多任务学习框架同时优化文档类型分类和重要度分级两个任务。", FONT_KAITI, FONT_KAITI, 11)

    # 4 实验
    section_heading(doc, "4　实验与分析", FONT_SIMHEI, FONT_SIMHEI, 12)
    body_para(doc, "实验使用真实政务数据集，包含15类共50,000份政务文档。", FONT_KAITI, FONT_KAITI, 11)
    body_para(doc, "实验结果表明，本文方法在Macro-F1和Micro-F1指标上分别达到94.3%和96.1%，较基线方法提升显著，验证了所提方法的有效性。", FONT_KAITI, FONT_KAITI, 11)

    # 5 结论
    section_heading(doc, "5　结论", FONT_SIMHEI, FONT_SIMHEI, 12)
    body_para(doc, "本文提出一种基于多任务学习与注意力机制的政务文档智能分类方法，在真实数据集上取得了优异的分类性能。未来工作将探索多模态政务文档（图文混合）的分类问题。", FONT_KAITI, FONT_KAITI, 11)

    # References heading
    section_heading(doc, "参考文献", FONT_SIMHEI, FONT_SIMHEI, 12)
    refs = [
        "[1] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding[C]//NAACL. 2019: 4171-4186.",
        "[2] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[J]. Advances in Neural Information Processing Systems, 2017: 5998-6008.",
        "[3] 王磊, 李强. 基于深度学习的政务文档自动分类研究[J]. 中文信息学报, 2025, 39(4): 112-120.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = FONT_LH
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(22)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    return doc

if __name__ == "__main__":
    base = "/home/zyu/.openclaw/workspaces/Main/docx-normalizer-ai/templates"

    doc = create_notification()
    doc.save(f"{base}/notification/notification-001-template.docx")
    print("Saved notification-001-template.docx")

    doc = create_report()
    doc.save(f"{base}/report/report-001-template.docx")
    print("Saved report-001-template.docx")

    doc = create_letter()
    doc.save(f"{base}/letter/letter-001-template.docx")
    print("Saved letter-001-template.docx")

    doc = create_academic()
    doc.save(f"{base}/academic/academic-001-template.docx")
    print("Saved academic-001-template.docx")
