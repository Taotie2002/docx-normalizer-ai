#!/usr/bin/env python3
"""
Apply Full - 完整Apply流程

处理流程:
    前置检查 → 备份（SHA-256）→ 结构签名 → 内容处理 → 验证 → 原子替换

F-12: SHA-256备份
F-13: 备份 + 回滚
完整性校验: 验证处理后文档无损坏
原子替换: 确保文档替换不出现中间状态

Usage:
    python scripts/apply_full.py <input_docx> <template_docx> <output_docx>
    python scripts/apply_full.py input.docx template.docx output.docx [--no-backup]

Example:
    python scripts/apply_full.py report.docx templates/government/gov-001-template.docx output.docx
"""

import argparse
import hashlib
import json
import os
import shutil
import sys
import tempfile
import uuid
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from scripts.dfgp_tool import 灌注公文
from scripts.backup import sha256_file, batch_backup


# ==========================================
# 完整性校验
# ==========================================

def verify_docx_integrity(path: str) -> dict:
    """
    验证DOCX文件完整性
    
    检查项:
    1. 文件可读取
    2. ZIP结构完整 (docx本质是zip)
    3. 包含必要的document.xml
    4. XML可解析
    """
    result = {
        "valid": False,
        "errors": [],
        "warnings": [],
        "sha256": None,
    }
    
    if not os.path.exists(path):
        result["errors"].append(f"文件不存在: {path}")
        return result
    
    # SHA-256
    result["sha256"] = sha256_file(path)
    
    # 尝试用python-docx打开
    try:
        from docx import Document
        doc = Document(path)
        para_count = len(doc.paragraphs)
        if para_count == 0:
            result["warnings"].append("文档无段落内容")
    except Exception as e:
        result["errors"].append(f"python-docx无法打开: {e}")
        return result
    
    # 检查ZIP结构
    try:
        import zipfile
        with zipfile.ZipFile(path, 'r') as zf:
            names = zf.namelist()
            # 检查必要文件
            required = ['word/document.xml']
            for req in required:
                if req not in names:
                    result["errors"].append(f"缺少必要文件: {req}")
            # 验证ZIP完整性
            bad_file = zf.testzip()
            if bad_file is not None:
                result["errors"].append(f"ZIP损坏: {bad_file}")
    except zipfile.BadZipFile as e:
        result["errors"].append(f"无效的ZIP格式: {e}")
        return result
    except Exception as e:
        result["errors"].append(f"ZIP验证失败: {e}")
        return result
    
    if not result["errors"]:
        result["valid"] = True
    
    return result


def compute_structure_signature(path: str) -> str:
    """
    计算文档结构签名
    用于验证处理前后结构一致性
    """
    from docx import Document
    from docx.oxml.ns import qn
    
    doc = Document(path)
    
    # 收集结构信息
    elements = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            # 获取格式信息
            alignment = str(para.alignment) if para.alignment else "None"
            elements.append({
                "i": i,
                "len": len(text),
                "align": alignment,
            })
    
    # 计算结构hash
    import json as json_mod
    structure_json = json_mod.dumps(elements, sort_keys=True)
    return hashlib.sha256(structure_json.encode()).hexdigest()[:16]


# ==========================================
# 原子替换
# ==========================================

def atomic_replace(src: str, dst: str) -> None:
    """
    原子替换文件
    使用rename确保不会出现中间状态
    """
    # 创建临时文件在同一目录下
    dst_dir = Path(dst).parent
    dst_dir.mkdir(parents=True, exist_ok=True)
    
    with tempfile.NamedTemporaryFile(dir=dst_dir, delete=False, suffix='.docx') as tmp:
        tmp_path = tmp.name
    
    try:
        # 复制到临时文件
        shutil.copy2(src, tmp_path)
        
        # 验证临时文件
        verify = verify_docx_integrity(tmp_path)
        if not verify["valid"]:
            raise RuntimeError(f"临时文件完整性校验失败: {verify['errors']}")
        
        # 原子替换
        os.replace(tmp_path, dst)
        
    except Exception:
        # 清理临时文件
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise


# ==========================================
# 主Apply流程
# ==========================================

def apply_full(
    input_path: str,
    template_path: str,
    output_path: str,
    backup_dir: str = "backup",
    skip_backup: bool = False,
) -> dict:
    """
    完整Apply流程
    
    Returns:
        执行报告 dict
    """
    report = {
        "success": False,
        "steps": [],
        "errors": [],
        "warnings": [],
    }
    
    def step(name: str, status: str, detail: str = ""):
        report["steps"].append({
            "name": name,
            "status": status,
            "detail": detail,
            "time": datetime.now().isoformat(timespec='milliseconds'),
        })
        status_icon = "✅" if status == "ok" else "❌" if status == "fail" else "⏳"
        print(f"  {status_icon} {name}: {detail}")
    
    input_path = os.path.abspath(input_path)
    output_path = os.path.abspath(output_path)
    template_path = os.path.abspath(template_path) if template_path else None
    
    print(f"\n{'='*60}")
    print(f"📋 Apply Full 流程开始")
    print(f"{'='*60}")
    print(f"  输入: {input_path}")
    print(f"  模板: {template_path}")
    print(f"  输出: {output_path}")
    
    # ---- Step 1: 前置检查 ----
    print(f"\n[1/6] 前置检查...")
    
    if not os.path.exists(input_path):
        step("前置检查", "fail", f"输入文件不存在: {input_path}")
        report["errors"].append(f"输入文件不存在: {input_path}")
        return report
    
    if template_path and not os.path.exists(template_path):
        step("前置检查", "fail", f"模板文件不存在: {template_path}")
        report["errors"].append(f"模板文件不存在: {template_path}")
        return report
    
    # 验证输入文件
    input_verify = verify_docx_integrity(input_path)
    if not input_verify["valid"]:
        step("前置检查", "fail", f"输入文件损坏: {input_verify['errors']}")
        report["errors"].extend(input_verify["errors"])
        return report
    
    step("前置检查", "ok", f"输入文件完整，SHA-256={input_verify['sha256'][:12]}...")
    
    # ---- Step 2: 备份 ----
    print(f"\n[2/6] 备份...")
    
    if skip_backup:
        step("备份", "skip", "已跳过")
        backup_entry = None
    else:
        try:
            manifest = batch_backup([input_path], backup_dir)
            backup_entry = manifest["backups"][-1] if manifest["backups"] else None
            if backup_entry:
                step("备份", "ok", f"备份到 {backup_entry['backup_filename']}，SHA-256={backup_entry['sha256'][:12]}...")
            else:
                step("备份", "fail", "备份失败")
                report["errors"].append("备份失败")
                return report
        except Exception as e:
            step("备份", "fail", str(e))
            report["errors"].append(f"备份异常: {e}")
            return report
    
    # ---- Step 3: 结构签名 ----
    print(f"\n[3/6] 结构签名...")
    
    try:
        pre_signature = compute_structure_signature(input_path)
        step("结构签名", "ok", f"预处理前签名={pre_signature}")
    except Exception as e:
        step("结构签名", "fail", str(e))
        report["errors"].append(f"结构签名失败: {e}")
        return report
    
    # ---- Step 4: 内容处理 ----
    print(f"\n[4/6] 内容处理...")
    
    # 生成临时输出文件
    tmp_output_dir = tempfile.mkdtemp(prefix="apply_full_")
    tmp_output = os.path.join(tmp_output_dir, Path(output_path).name)
    
    try:
        灌注公文(input_path, tmp_output, template_path)
        step("内容处理", "ok", f"灌注完成: {tmp_output}")
    except Exception as e:
        step("内容处理", "fail", str(e))
        report["errors"].append(f"内容处理失败: {e}")
        # 清理临时目录
        shutil.rmtree(tmp_output_dir, ignore_errors=True)
        return report
    
    # ---- Step 5: 验证 ----
    print(f"\n[5/6] 完整性验证...")
    
    output_verify = verify_docx_integrity(tmp_output)
    if not output_verify["valid"]:
        step("完整性验证", "fail", f"输出文件损坏: {output_verify['errors']}")
        report["errors"].extend(output_verify["errors"])
        shutil.rmtree(tmp_output_dir, ignore_errors=True)
        return report
    
    step("完整性验证", "ok", f"输出文件完整，SHA-256={output_verify['sha256'][:12]}...")
    
    # 结构签名验证（可选）
    try:
        post_signature = compute_structure_signature(tmp_output)
        step("结构签名(后)", "ok", f"处理后签名={post_signature}")
        if pre_signature != post_signature:
            report["warnings"].append(f"结构发生变化: {pre_signature} → {post_signature}")
    except Exception as e:
        step("结构签名(后)", "warn", str(e))
        report["warnings"].append(f"后结构签名失败: {e}")
    
    # ---- Step 6: 原子替换 ----
    print(f"\n[6/6] 原子替换...")
    
    try:
        atomic_replace(tmp_output, output_path)
        step("原子替换", "ok", f"替换完成: {output_path}")
    except Exception as e:
        step("原子替换", "fail", str(e))
        report["errors"].append(f"原子替换失败: {e}")
        shutil.rmtree(tmp_output_dir, ignore_errors=True)
        return report
    
    # 清理临时目录
    shutil.rmtree(tmp_output_dir, ignore_errors=True)
    
    # ---- 最终校验 ----
    print(f"\n[完成] 最终校验...")
    
    final_verify = verify_docx_integrity(output_path)
    if not final_verify["valid"]:
        step("最终校验", "fail", f"输出文件损坏: {final_verify['errors']}")
        report["errors"].extend(final_verify["errors"])
        return report
    
    # SHA-256对比
    if backup_entry:
        final_sha = final_verify["sha256"]
        if final_sha == backup_entry["sha256"]:
            step("最终校验", "warn", "输出与备份完全相同，可能未做任何修改")
            report["warnings"].append("输出与备份完全相同")
        else:
            step("最终校验", "ok", f"输出SHA-256={final_sha[:12]}...，与备份不同")
    
    report["success"] = True
    report["output_sha256"] = final_verify["sha256"]
    
    print(f"\n{'='*60}")
    print(f"✅ Apply Full 流程完成!")
    print(f"{'='*60}")
    
    return report


def main():
    parser = argparse.ArgumentParser(
        description="完整Apply流程 - 备份 + 灌注 + 验证 + 原子替换",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    python scripts/apply_full.py report.docx templates/government/gov-001-template.docx output.docx
    python scripts/apply_full.py report.docx templates/government/gov-001-template.docx output.docx --no-backup
    python scripts/apply_full.py report.docx templates/government/gov-001-template.docx output.docx --backup-dir my_backups
        """
    )
    parser.add_argument("input", help="输入文档路径")
    parser.add_argument("template", help="模板文档路径")
    parser.add_argument("output", help="输出文档路径")
    parser.add_argument("--backup-dir", "-d", default="backup", help="备份目录 (默认: backup/)")
    parser.add_argument("--no-backup", action='store_true', help="跳过备份步骤")
    parser.add_argument("--report", "-r", action='store_true', help="输出JSON格式报告")
    args = parser.parse_args()

    report = apply_full(
        args.input,
        args.template,
        args.output,
        backup_dir=args.backup_dir,
        skip_backup=args.no_backup,
    )

    if args.report:
        print(json.dumps(report, ensure_ascii=False, indent=2))

    sys.exit(0 if report["success"] else 1)


if __name__ == "__main__":
    main()
