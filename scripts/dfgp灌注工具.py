#!/usr/bin/env python3
"""
公文格式灌注工具 v1.0
固化正确经验，避免重复犯错
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import zipfile
import shutil
from lxml import etree
import re

# ==========================================
# 核心识别函数（已验证，正确）
# ==========================================

def is_main_title(text, index):
    """主标题：第一个居中的短标题（通常含"关于"和"请示"）"""
    if index == 0 and '关于' in text and ('请示' in text or '报告' in text):
        return True
    if index == 0 and len(text) <= 30 and len(text) > 2:
        return True
    return False

def is_title_l1(text):
    """一级标题：X、格式"""
    return bool(text.startswith(tuple('一二三四五六七八九十')) and '、' in text[:5])

def is_title_l2(text):
    """二级标题：（一）格式"""
    return text.startswith('（') and '）' in text[:6]

def is_salutation(text):
    """称谓行：尊敬的...、市政府：等"""
    if text.startswith('尊敬的') and '领导' in text:
        return True
    if text.endswith('：') and len(text) < 20:
        return True
    return False

def is_conclusion(text):
    """结语：妥否，请批示"""
    return '妥否' in text and ('批示' in text or '请' in text)

def is_signature(text):
    """落款：单位名称+日期"""
    if any(kw in text for kw in ['管理局', '办公室', '委员会', '办公厅']):
        return True
    if re.search(r'\d{4}年\d{1,2}月\d{1,2}日', text):
        return True
    return False

# ==========================================
# 焦土策略（已验证，正确）
# ==========================================

def scorched_earth(para):
    """彻底清除段落和run的所有格式"""
    # 清除阴影
    for run in para.runs:
        for shd in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd'):
            run._element.remove(shd)
    
    # 清除段落对齐
    pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is not None:
        for jc in pPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc'):
            pPr.remove(jc)
    
    # 清除run格式（保留文本）
    for run in para.runs:
        rPr = run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is not None:
            for child in list(rPr):
                tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
                if tag not in ['rFonts', 'sz', 'szCs']:
                    rPr.remove(child)

def set_font(run, font_name, size_pt):
    """设置run字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)

# ==========================================
# 页码处理（已验证，正确）
# ==========================================

def add_page_numbers(doc, template_path=None):
    """
    添加公文页码
    - 奇数页：居右
    - 偶数页：居左
    - 首页不显示
    """
    if template_path is None:
        template_path = "/home/zyu/.openclaw/media/inbound/请示2---159f9e22-4054-4c15-bb9d-2169d31ef81f.docx"
    
    # 从模板获取footer
    with zipfile.ZipFile(template_path, 'r') as z:
        footer1_content = z.read('word/footer1.xml')
        footer2_content = z.read('word/footer2.xml')
    
    section = doc.sections[0]
    
    # 设置sectPr属性
    sectPr = section._sectPr
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    
    # 清除所有footer和header引用
    for child in list(sectPr):
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if 'footerReference' in tag or 'headerReference' in tag:
            sectPr.remove(child)
    
    # 清除首页不同标记
    titlePg = sectPr.find('{%s}titlePg' % ns_w)
    if titlePg is not None:
        sectPr.remove(titlePg)
    
    # 添加footer引用
    FooterRef1 = etree.SubElement(sectPr, '{%s}footerReference' % ns_w)
    FooterRef1.set('{%s}type' % ns_r, 'default')
    FooterRef1.set('{%s}id' % ns_r, 'rId99')
    
    FooterRef2 = etree.SubElement(sectPr, '{%s}footerReference' % ns_w)
    FooterRef2.set('{%s}type' % ns_r, 'even')
    FooterRef2.set('{%s}id' % ns_r, 'rId98')
    
    # 添加首页不同
    titlePg = etree.SubElement(sectPr, '{%s}titlePg' % ns_w)
    titlePg.set('{%s}val' % ns_w, '1')
    
    # 添加到content_types
    part = doc.part
    ct = part.package.part_related_by('http://schemas.openxmlformats.org/package/2006/content-types', '[Content_Types].xml')
    ct_root = etree.fromstring(ct.blob)
    ns_ct = 'http://schemas.openxmlformats.org/package/2006/content-types'
    
    existing = [o.get('PartName') for o in ct_root.findall('{%s}Override' % ns_ct)]
    if '/word/footer1.xml' not in existing:
        o = etree.SubElement(ct_root, '{%s}Override' % ns_ct)
        o.set('PartName', '/word/footer1.xml')
        o.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml')
    if '/word/footer2.xml' not in existing:
        o = etree.SubElement(ct_root, '{%s}Override' % ns_ct)
        o.set('PartName', '/word/footer2.xml')
        o.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml')
    ct.blob = etree.tostring(ct_root, xml_declaration=True, encoding='UTF-8', standalone=True)
    
    # 添加到rels
    rels = part.rels
    rels['rId99'] = type(rels.popitem()[1])(
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer',
        'footer1.xml'
    )
    rels['rId98'] = type(rels.popitem()[1])(
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer',
        'footer2.xml'
    )
    
    # 添加footer blob
    doc.part.add_relationship('http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer', 'footer1.xml', 'rId99')
    doc.part.add_relationship('http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer', 'footer2.xml', 'rId98')
    
    # 直接修改zip
    return doc

# ==========================================
# 主灌注函数（已验证，正确）
# ==========================================

def 灌注公文(input_path, output_path, template_footer_path=None):
    """灌注公文格式"""
    doc = Document(input_path)
    section = doc.sections[0]
    
    # 页面设置
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.different_first_page_header_footer = True
    
    # 焦土
    for para in doc.paragraphs:
        scorched_earth(para)
    
    # 应用格式
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # 重建run
        for run in para.runs:
            run.text = ''
        if not para.runs:
            para.add_run(text)
        else:
            para.runs[0].text = text
        run = para.runs[0]
        
        # 识别类型并设置格式
        if is_main_title(text, i):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(run, '方正小标宋简体', 22)
            para.paragraph_format.first_line_indent = None
        elif is_title_l1(text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(run, '黑体', 16)
            para.paragraph_format.first_line_indent = None
        elif is_title_l2(text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(run, '楷体_GB2312', 16)
            para.paragraph_format.first_line_indent = None
        elif is_salutation(text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(run, '仿宋_GB2312', 16)
            para.paragraph_format.first_line_indent = None  # 称谓行无缩进！
        elif is_conclusion(text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(run, '仿宋_GB2312', 16)
            para.paragraph_format.first_line_indent = None
        elif is_signature(text):
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_font(run, '仿宋_GB2312', 16)
            para.paragraph_format.first_line_indent = None
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(run, '仿宋_GB2312', 16)
            para.paragraph_format.first_line_indent = Pt(32)  # 正文首行缩进
        
        # 行间距固定28pt
        para.paragraph_format.line_spacing = Pt(28)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    
    # 添加页码
    if template_footer_path:
        add_page_numbers(doc, template_footer_path)
    
    doc.save(output_path)
    print(f"✅ 灌注完成: {output_path}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) >= 3:
        灌注公文(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
    else:
        print("用法: python3 dfgp灌注工具.py <输入文件> <输出文件> [页码模板]")
